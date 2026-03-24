
from __future__ import annotations
from typing import List
from pathlib import Path
import re
import bibtexparser
from .models import BibEntry

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


ENTRY_START_RE = re.compile(r"@\w+\s*(\{|[^\s])", re.IGNORECASE)
REFERENCE_HEADING_RE = re.compile(r"^\s*(?:\d+\.\s*)?(references|bibliography|works cited)\s*$", re.IGNORECASE | re.MULTILINE)
DOI_RE = re.compile(r"(?:https?://(?:dx\.)?doi\.org/|doi\s*[:]?\s*)(10\.\d{4,9}/[-._;()/:A-Z0-9]+)", re.IGNORECASE)
YEAR_RE = re.compile(r"(?<!\d)(19\d{2}|20\d{2})(?!\d)")
NUMBERED_REF_RE = re.compile(r"^\s*(?:\[(\d+)\]|(\d+)\.|\((\d+)\))\s+")
ARXIV_RE = re.compile(r"\barXiv\s*:?\s*([A-Za-z.-]+/\d{7}|\d{4}\.\d{4,5})(?:v\d+)?", re.IGNORECASE)


def _extract_bibtex_text(raw_text: str) -> str:
    if "@" not in raw_text:
        raise ValueError("No BibTeX entries found in uploaded file.")
    match = ENTRY_START_RE.search(raw_text)
    if not match:
        raise ValueError("No BibTeX entries found in uploaded file.")
    return raw_text[match.start():]


def _split_entries(bib_text: str) -> list[str]:
    starts = list(ENTRY_START_RE.finditer(bib_text))
    chunks = []
    for i, m in enumerate(starts):
        start = m.start()
        end = starts[i + 1].start() if i + 1 < len(starts) else len(bib_text)
        chunks.append(bib_text[start:end].strip())
    return chunks


def _parse_entry_chunk(chunk: str):
    try:
        db = bibtexparser.loads(chunk)
        if db.entries:
            return db.entries[0]
    except Exception:
        pass
    return None


def _recover_entry_chunk(chunk: str):
    header = re.match(r"@(\w+)\s*(?:\{|\s)\s*([^,\s\{]+)", chunk, re.IGNORECASE)
    if not header:
        return None
    entry_type, entry_key = header.group(1).lower(), header.group(2).strip()
    fields = {}
    for line in chunk.splitlines()[1:]:
        line = line.strip().rstrip(",")
        if not line or line == "}":
            continue
        if "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip(",").strip()
        if value.startswith("{") and value.endswith("}"):
            value = value[1:-1]
        elif value.startswith('"') and value.endswith('"'):
            value = value[1:-1]
        fields[key] = value
    return {"ID": entry_key, "ENTRYTYPE": entry_type, **fields}


def _read_pdf_text(path: Path) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF is required for PDF ingestion.")
    out = []
    with fitz.open(path) as doc:
        for page in doc:
            out.append(page.get_text("text"))
    return "\n".join(out)




def _read_pdf_references_text(path: Path) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF is required for PDF ingestion.")
    pages = []
    with fitz.open(path) as doc:
        for page in doc:
            pages.append(page.get_text("text"))
    # Prefer pages from the last page containing a references heading onward.
    ref_page_idx = None
    for i, page_text in enumerate(pages):
        if REFERENCE_HEADING_RE.search(page_text):
            ref_page_idx = i
    if ref_page_idx is not None:
        joined = "\n".join(pages[ref_page_idx:])
        return _extract_references_section(joined)
    return _extract_references_section("\n".join(pages))

def _read_docx_text(path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX ingestion.")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def _read_docx_paragraphs(path: Path) -> list[str]:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX ingestion.")
    doc = Document(path)
    return [p.text for p in doc.paragraphs]


def _extract_references_paragraphs(paragraphs: list[str]) -> list[str]:
    cleaned = [p.strip() for p in paragraphs]
    ref_idx = None
    for i in range(len(cleaned) - 1, -1, -1):
        if REFERENCE_HEADING_RE.match(cleaned[i] or ""):
            ref_idx = i
            break
    if ref_idx is not None:
        tail = cleaned[ref_idx + 1:]
    else:
        tail = cleaned[max(0, len(cleaned) // 2):]
    return [p for p in tail if p]


def _looks_like_reference_paragraph(paragraph: str) -> bool:
    if len(paragraph.split()) < 4:
        return False
    if not YEAR_RE.search(paragraph):
        return False
    start = paragraph[:140]
    if re.match(r"^[A-ZÀ-ÖØ-Þ][\w'’`.-]+(?:\s+[A-ZÀ-ÖØ-Þ][\w'’`.-]+)*,", start):
        return True
    if re.match(r"^[A-ZÀ-ÖØ-Þ][\w'’`.-]+\s+[A-Z]\.", start):
        return True
    return False




def _preprocess_pdf_text(text: str) -> str:
    # Join words broken by PDF line-wrap hyphenation
    text = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", text)
    # Remove soft / odd hyphenation artifacts and normalize dash-like chars
    text = text.replace("￾", "")
    text = text.replace("­", "")
    for ch in ["–", "—", "−"]:
        text = text.replace(ch, "-")
    # Normalize whitespace around line breaks before later splitting
    text = re.sub(r"[ 	]+", " ", text)
    return text


def _preprocess_reference_section(ref_text: str) -> str:
    ref_text = _preprocess_pdf_text(ref_text)
    # Merge wrapped lines inside references but preserve numbered starts like [1], 1., (1)
    ref_text = re.sub(r"\n(?!\s*(?:\[\d+\]|\d+\.|\(\d+\)))", " ", ref_text)
    ref_text = re.sub(r"\s+", " ", ref_text).strip()
    # Reintroduce line breaks before numbered references for easier splitting
    ref_text = re.sub(r"\s+(?=(?:\[\d+\]|\d+\.|\(\d+\))\s)", "\n", ref_text)
    return ref_text

def _extract_references_section(text: str) -> str:
    matches = list(REFERENCE_HEADING_RE.finditer(text))
    if matches:
        m = matches[-1]
        return text[m.end():].strip()
    # Fallback: take the tail if a heading is not present.
    lines = [ln.rstrip() for ln in text.splitlines()]
    if len(lines) > 30:
        return "\n".join(lines[(2 * len(lines))//3:]).strip()
    return text.strip()


def _split_reference_blocks(ref_text: str) -> list[str]:
    lines = [ln.rstrip() for ln in ref_text.splitlines()]
    blocks: list[list[str]] = []
    current: list[str] = []
    saw_numbered = False
    for raw in lines:
        line = raw.strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        if NUMBERED_REF_RE.match(line):
            saw_numbered = True
            if current:
                blocks.append(current)
            current = [NUMBERED_REF_RE.sub("", line, count=1).strip()]
            continue
        if saw_numbered:
            current.append(line)
        else:
            if current:
                current.append(line)
            else:
                current = [line]
    if current:
        blocks.append(current)
    merged = [" ".join(b).strip() for b in blocks if " ".join(b).strip()]
    merged = [b for b in merged if len(b.split()) >= 4]
    if len(merged) == 1:
        giant = merged[0]
        year_positions = [m.start() for m in re.finditer(r"\(?(?:19\d{2}|20\d{2})\)?", giant)]
        if len(year_positions) >= 3:
            paras = [ln.strip() for ln in ref_text.splitlines() if ln.strip()]
            para_like = [p for p in paras if _looks_like_reference_paragraph(p)]
            if len(para_like) >= 3:
                return para_like
            # fallback: split on likely author-year starts inside the giant block
            starts = [m.start() for m in re.finditer(r"(?=(?:^|\s)([A-ZÀ-ÖØ-Þ][\w'’`.-]+(?:,\s*[A-Z](?:\.[A-Z])?\.?)+(?:,?\s*(?:&|and)\s*[A-ZÀ-ÖØ-Þ][\w'’`.-]+(?:,\s*[A-Z](?:\.[A-Z])?\.?)+)*\s*\((?:19\d{2}|20\d{2})\)))", giant)]
            if len(starts) >= 2:
                chunks=[]
                for i,s in enumerate(starts):
                    end = starts[i+1] if i+1 < len(starts) else len(giant)
                    chunk = giant[s:end].strip()
                    if len(chunk.split()) >= 4:
                        chunks.append(chunk)
                if len(chunks) >= 2:
                    return chunks
    return merged


def _clean_reference_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s*doi\s*:\s*", " doi:", text, flags=re.IGNORECASE)
    return text


def _extract_title_from_reference(text: str, year_match: re.Match | None) -> str:
    quoted = re.search(r'["“](.+?)["”]', text)
    if quoted:
        return quoted.group(1).strip().strip('.')
    search_text = text
    if year_match:
        search_text = text[year_match.end():].lstrip(").,; ")
    parts = [p.strip(" .;") for p in re.split(r"\.\s+", search_text) if p.strip(" .;")]
    for part in parts:
        low = part.lower()
        if low.startswith(("doi:", "https://doi.org/", "arxiv", "in ")):
            continue
        if 2 <= len(part.split()) <= 25:
            return part
    return ""


def _extract_author_from_reference(text: str, year_match: re.Match | None) -> str:
    if not year_match:
        return ""
    author_part = text[:year_match.start()].strip(" ,.;()")
    author_part = re.sub(r"\bet al\.?$", "", author_part, flags=re.IGNORECASE).strip(" ,.;")
    return author_part


def _reference_block_to_entry(block: str, index: int) -> BibEntry:
    text = _clean_reference_text(block)
    doi_match = DOI_RE.search(text)
    year_match = YEAR_RE.search(text)
    arxiv_match = ARXIV_RE.search(text)
    title = _extract_title_from_reference(text, year_match)
    author = _extract_author_from_reference(text, year_match)
    year = year_match.group(1) if year_match else ""
    fields = {}
    if title:
        fields["title"] = title
    if author:
        fields["author"] = author
    if year:
        fields["year"] = year
    if doi_match:
        fields["doi"] = doi_match.group(1).rstrip('.,;')
    if arxiv_match:
        fields["eprint"] = arxiv_match.group(1)
        fields["archivePrefix"] = "arXiv"
    entry_type = "article" if doi_match else ("misc" if arxiv_match else "article")
    raw = {"ID": f"ref_{index:03d}", "ENTRYTYPE": entry_type, **fields, "_source_reference": text}
    return BibEntry(entry_key=f"ref_{index:03d}", entry_type=entry_type, fields=fields, raw_entry=raw)


def load_bib_entries_from_document_text(raw_text: str) -> List[BibEntry]:
    ref_text = _preprocess_reference_section(_extract_references_section(raw_text))
    blocks = _split_reference_blocks(ref_text)
    entries = []
    for i, block in enumerate(blocks, start=1):
        entries.append(_reference_block_to_entry(block, i))
    if not entries:
        raise ValueError("No reference entries detected in document.")
    return entries


def load_bib_entries_from_text(raw_text: str) -> List[BibEntry]:
    if "@" not in raw_text:
        return load_bib_entries_from_document_text(raw_text)
    bib_text = _extract_bibtex_text(raw_text)
    entries = []
    for chunk in _split_entries(bib_text):
        raw = _parse_entry_chunk(chunk)
        if raw is None:
            raw = _recover_entry_chunk(chunk)
        if raw is None:
            continue
        entry_key = raw.get("ID", "")
        entry_type = raw.get("ENTRYTYPE", "misc")
        fields = {k: v for k, v in raw.items() if k not in {"ID", "ENTRYTYPE"}}
        entries.append(BibEntry(entry_key=entry_key, entry_type=entry_type, fields=fields, raw_entry=raw))
    return entries


def load_bib_entries_from_docx(path: str) -> List[BibEntry]:
    paragraphs = _read_docx_paragraphs(Path(path))
    ref_paragraphs = _extract_references_paragraphs(paragraphs)
    blocks = [p.strip() for p in ref_paragraphs if _looks_like_reference_paragraph(p)]
    if not blocks:
        return load_bib_entries_from_document_text("\n".join(ref_paragraphs))
    return [_reference_block_to_entry(block, i) for i, block in enumerate(blocks, start=1)]


def load_bib_entries(path: str) -> List[BibEntry]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        ref_text = _preprocess_reference_section(_read_pdf_references_text(p))
        blocks = _split_reference_blocks(ref_text)
        if not blocks:
            raise ValueError("No reference entries detected in PDF.")
        return [_reference_block_to_entry(block, i) for i, block in enumerate(blocks, start=1)]
    if suffix == ".docx":
        return load_bib_entries_from_docx(str(p))
    raw_text = p.read_text(encoding="utf-8")
    return load_bib_entries_from_text(raw_text)


def dump_bib_entries(entries: List[BibEntry], path: str) -> None:
    db = bibtexparser.bibdatabase.BibDatabase()
    db.entries = []
    for e in entries:
        raw = {"ID": e.entry_key, "ENTRYTYPE": e.entry_type}
        raw.update(e.fields)
        db.entries.append(raw)
    writer = bibtexparser.bwriter.BibTexWriter()
    writer.indent = "  "
    with open(path, "w", encoding="utf-8") as f:
        f.write(writer.write(db))
